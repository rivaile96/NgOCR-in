#!/usr/bin/env python3
import os
import sys
import pytesseract
import cv2
from termcolor import colored, cprint
from deep_translator import GoogleTranslator
from docx import Document
from langdetect import detect
from datetime import datetime

# Warna untuk terminal
RED = '\033[91m'
GREEN = '\033[92m'
YELLOW = '\033[93m'
BLUE = '\033[94m'
MAGENTA = '\033[95m'
CYAN = '\033[96m'
WHITE = '\033[97m'
RESET = '\033[0m'

ascii_art = f"""{CYAN}
 ███▄    █   ▄████  ▒█████   ▄████▄   ██▀███   ██▓ ███▄    █ 
 ██ ▀█   █  ██▒ ▀█▒▒██▒  ██▒▒██▀ ▀█  ▓██ ▒ ██▒▓██▒ ██ ▀█   █ 
▓██  ▀█ ██▒▒██░▄▄▄░▒██░  ██▒▒▓█    ▄ ▓██ ░▄█ ▒▒██▒▓██  ▀█ ██▒
▓██▒  ▐▌██▒░▓█  ██▓▒██   ██░▒▓▓▄ ▄██▒▒██▀▀█▄  ░██░▓██▒  ▐▌██▒
▒██░   ▓██░░▒▓███▀▒░ ████▓▒░▒ ▓███▀ ░░██▓ ▒██▒░██░▒██░   ▓██░
░ ▒░   ▒ ▒  ░▒   ▒ ░ ▒░▒░▒░ ░ ░▒ ▒  ░░ ▒▓ ░▒▓░░▓  ░ ▒░   ▒ ▒ 
░ ░░   ░ ▒░  ░   ░   ░ ▒ ▒░   ░  ▒     ░▒ ░ ▒░ ▒ ░░ ░░   ░ ▒░
   ░   ░ ░ ░ ░   ░ ░ ░ ░ ▒  ░          ░░   ░  ▒ ░   ░   ░ ░ 
         ░       ░     ░ ░  ░ ░         ░      ░           ░ 
                            ░                                 
{YELLOW}                 Author: Rivaile
                 Version: v2.0{RESET}
"""

def ocr_image(image_path, lang, psm):
    custom_config = f"--psm {psm}" if psm else ""
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    text = pytesseract.image_to_string(gray, lang=lang, config=custom_config)
    return text

def translate_text(text, target_lang):
    try:
        translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
        return translated
    except Exception as e:
        return f"{RED}Gagal menerjemahkan: {e}{RESET}"

def save_results(original_text, translated_text):
    with open("ocr_result.txt", "w") as f:
        f.write("=== HASIL OCR ===\n\n")
        f.write(original_text + "\n\n")
        if translated_text:
            f.write("=== HASIL TERJEMAHAN ===\n\n")
            f.write(translated_text)

    doc = Document()
    doc.add_heading("Hasil OCR", level=1)
    doc.add_paragraph(original_text)
    if translated_text:
        doc.add_heading("Hasil Terjemahan", level=1)
        doc.add_paragraph(translated_text)
    doc.save("ocr_result.docx")

def detect_language(text):
    try:
        return detect(text)
    except:
        return "unknown"

def ocr_process(image_path, lang_code):
    psm = input("Masukkan PSM (default 6): ") or "6"
    print(f"\n⏳ Menjalankan OCR...\n")
    result = ocr_image(image_path, lang_code, psm)
    print(f"{GREEN}=== HASIL OCR ==={RESET}\n")
    print(result)

    translate = input("\nTerjemahkan hasil? (y/n): ").lower()
    translated = ""
    if translate == 'y':
        target_lang = input("Masukkan kode bahasa target (cth: id, en, ja, fr): ")
        print(f"\n{YELLOW}=== HASIL TERJEMAHAN ==={RESET}\n")
        translated = translate_text(result, target_lang)
        print(f"{YELLOW}{translated}{RESET}")

    save = input("\nSimpan hasil? (y/n): ").lower()
    if save == 'y':
        save_results(result, translated)
        print(f"\n{GREEN}✔ Hasil berhasil disimpan sebagai 'ocr_result.txt' dan 'ocr_result.docx'{RESET}")

def pilih_bahasa():
    print("\n🗣 Pilih Bahasa OCR:")
    daftar_bahasa = {
        "1": ("eng", "English"),
        "2": ("ind", "Indonesian"),
        "3": ("chi_sim", "Chinese (Simplified)"),
        "4": ("chi_tra", "Chinese (Traditional)"),
        "5": ("jpn", "Japanese"),
        "6": ("kor", "Korean"),
        "7": ("ara", "Arabic"),
        "8": ("rus", "Russian"),
        "9": ("auto", "Auto-Detect")
    }
    for key, (_, name) in daftar_bahasa.items():
        print(f"{key}. {name}")
    pilihan = input("Pilih (1-9): ")
    return daftar_bahasa.get(pilihan, ("eng", "English"))[0]

def menu():
    print(ascii_art)
    print(f"{MAGENTA}🔸 Menu Utama 🔸{RESET}")
    print(f"{WHITE}1. {BLUE}OCR Gambar Tunggal\n{WHITE}2. {BLUE}OCR Folder (Batch)\n{WHITE}99. {RED}Keluar{RESET}")
    pilihan = input("\nMasukkan pilihan: ")

    if pilihan == "1":
        lang = pilih_bahasa()
        path = input("\n📁 Masukkan path gambar: ")
        if os.path.isfile(path):
            ocr_process(path, lang)
        else:
            print(f"{RED}File tidak ditemukan.{RESET}")
    elif pilihan == "2":
        lang = pilih_bahasa()
        folder = input("\n📁 Masukkan path folder: ")
        if os.path.isdir(folder):
            for filename in os.listdir(folder):
                if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                    print(f"\n{CYAN}📄 Memproses: {filename}{RESET}")
                    path = os.path.join(folder, filename)
                    ocr_process(path, lang)
        else:
            print(f"{RED}Folder tidak ditemukan.{RESET}")
    elif pilihan == "99":
        print(f"{GREEN}Sampai jumpa!{RESET}")
        sys.exit()
    else:
        print(f"{RED}Pilihan tidak valid.{RESET}")

if __name__ == "__main__":
    menu()
